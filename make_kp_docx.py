from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = r"c:\Users\freak\TRAE\TechnoPan\КП_TechnoPan_XTeamPro.docx"

# ── palette ──────────────────────────────────────────────────────────────────
DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # почти чёрный синий
ACCENT = RGBColor(0x16, 0x21, 0x3E)   # тёмно-синий
BLUE   = RGBColor(0x0F, 0x3C, 0x96)   # основной синий
LIGHT  = RGBColor(0xF0, 0xF4, 0xFF)   # светло-голубой фон
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0x55, 0x55, 0x55)
LINE   = RGBColor(0xCC, 0xD6, 0xF0)


def rgb_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(rgb))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """kwargs: top, bottom, left, right — each dict(sz, color, val)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        opts = kwargs.get(side)
        if opts:
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), opts.get("val", "single"))
            border.set(qn("w:sz"), str(opts.get("sz", 4)))
            border.set(qn("w:color"), opts.get("color", "auto"))
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, size=11,
            color: RGBColor = None, font="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    run.font.color.rgb = color or DARK
    return run


def heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    para.paragraph_format.space_after  = Pt(6)
    size = 16 if level == 1 else 13
    color = BLUE if level == 1 else ACCENT
    add_run(para, text, bold=True, size=size, color=color)
    # нижняя линия под h1
    if level == 1:
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:color"), rgb_hex(BLUE))
        pBdr.append(bottom)
        pPr.append(pBdr)
    return para


def body(doc, text, space_after=6, color=None, size=11):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(space_after)
    add_run(para, text, size=size, color=color or DARK)
    return para


def bullet(doc, text, bold_part=None):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.left_indent  = Cm(0.8)
    if bold_part and text.startswith(bold_part):
        add_run(para, bold_part, bold=True, size=11)
        add_run(para, text[len(bold_part):], size=11)
    else:
        add_run(para, text, size=11)
    return para


def add_table(doc, headers, rows, col_widths=None,
              header_bg=BLUE, header_fg=WHITE,
              alt_bg=LIGHT):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"

    border_opts = dict(sz=4, color=rgb_hex(LINE), val="single")

    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        add_run(p, h, bold=True, size=10, color=header_fg)

    # data rows
    for r_idx, row_data in enumerate(rows):
        bg = alt_bg if r_idx % 2 == 0 else WHITE
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cell = cells[c_idx]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            bold = isinstance(val, str) and val.startswith("**") and val.endswith("**")
            text = val.strip("*") if bold else str(val)
            add_run(p, text, bold=bold, size=10)

    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def callout_box(doc, text, bg=LIGHT, accent=BLUE):
    """Однострочный акцентный блок через таблицу 1×1."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    # левая синяя граница
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "right", "bottom"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tcBorders.append(b)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), rgb_hex(accent))
    tcBorders.append(left)
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.4)
    add_run(p, text, size=11, italic=True, color=ACCENT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def stage_block(doc, title, duration, price, items):
    """Блок этапа с шапкой-таблицей и буллетами."""
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    bg_dark = ACCENT
    widths = [9, 3, 4]
    data = [title, duration, price]
    bolds = [True, False, True]
    for i, (val, w, b) in enumerate(zip(data, widths, bolds)):
        cell = table.cell(0, i)
        set_cell_bg(cell, bg_dark)
        cell.width = Cm(w)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.left_indent  = Cm(0.3) if i == 0 else Cm(0)
        add_run(p, val, bold=b, size=11, color=WHITE)

    for item in items:
        bullet(doc, item)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ═════════════════════════════════════════════════════════════════════════════
doc = Document()

# Поля страницы
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

# ── ШАПКА ────────────────────────────────────────────────────────────────────
header_table = doc.add_table(rows=1, cols=2)
header_table.style = "Table Grid"
left_cell  = header_table.cell(0, 0)
right_cell = header_table.cell(0, 1)
set_cell_bg(left_cell,  ACCENT)
set_cell_bg(right_cell, DARK)
left_cell.width  = Cm(11)
right_cell.width = Cm(6)

pl = left_cell.paragraphs[0]
pl.paragraph_format.space_before = Pt(10)
pl.paragraph_format.space_after  = Pt(2)
pl.paragraph_format.left_indent  = Cm(0.5)
add_run(pl, "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ", bold=True, size=18, color=WHITE)

pl2 = left_cell.add_paragraph()
pl2.paragraph_format.space_before = Pt(0)
pl2.paragraph_format.space_after  = Pt(10)
pl2.paragraph_format.left_indent  = Cm(0.5)
add_run(pl2, "XTeam.Pro  →  TechnoPan", size=11, color=RGBColor(0xAA, 0xBB, 0xEE))

pr = right_cell.paragraphs[0]
pr.paragraph_format.space_before = Pt(12)
pr.paragraph_format.space_after  = Pt(2)
pr.paragraph_format.left_indent  = Cm(0.4)
add_run(pr, "Автоматизация формирования\nспецификаций из DWG-файлов\nраскладки панелей",
        size=10, color=WHITE, italic=True)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ── 1. МЫ ИЗУЧИЛИ ВАШИ ФАЙЛЫ ─────────────────────────────────────────────────
heading(doc, "Мы изучили ваши реальные файлы")

body(doc, "Нам были переданы три реальных проекта TechnoPan:")

add_table(doc,
    headers=["Объект", "Файл раскладки", "Файл спецификации", "Разрыв"],
    rows=[
        ["Магазин в Чулыме",          "01.09.2025", "16.09.2025", "**15 дней**"],
        ["Склад готовой продукции",    "15.12.2025", "16.12.2025", "1 день"],
        ["Цех, перегородки АБК",       "02.12.2025", "04.12.2025", "2 дня"],
    ],
    col_widths=[6.5, 3.5, 3.5, 2.5],
)

callout_box(doc,
    "Каждый день между готовым чертежом и отправленной спецификацией — это замороженные деньги, "
    "незакрытая сделка и риск, что клиент за это время получит предложение от конкурента. "
    "При 15-дневном разрыве, как по «Магазину в Чулыме», это уже не вопрос удобства — "
    "это вопрос выживаемости сделки."
)

p = doc.add_paragraph()
add_run(p, "Конкретный пример — «Склад готовой продукции»:", bold=True, size=11)

for item in [
    "189 панелей, 36 уникальных типоразмеров",
    "Два типа: кровельные (п) и стеновые (с), ширина 1000 мм / стандартная",
    "Длины от 1130 до 7540 мм, самая частая позиция — п 5980 мм: 51 шт.",
    "Всё это нужно пересчитать вручную, сгруппировать и занести в таблицу",
]:
    bullet(doc, item)

body(doc, "Именно эту работу мы автоматизируем.", color=BLUE, size=11)

# ── 2. КАК РАБОТАЕТ СИСТЕМА ───────────────────────────────────────────────────
heading(doc, "Как работает система")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
add_run(p, "Вход: ", bold=True, size=11)
add_run(p, "DWG-файл раскладки панелей (как есть, без переделки чертежей)", size=11)
p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(6)
add_run(p2, "Выход: ", bold=True, size=11)
add_run(p2, "готовая спецификация Excel — за 10–30 секунд", size=11)

body(doc,
    "Система читает ваши DWG напрямую, распознаёт маркировки панелей по слоям "
    "(«Нумерация», «Нумерация 1000»), считает количество панелей каждого типоразмера "
    "и формирует сводную таблицу в вашем корпоративном шаблоне."
)
body(doc,
    "Ничего менять в процессе проектирования не нужно — система адаптируется под ваши файлы, а не наоборот.",
    color=BLUE
)

# ── 3. ЧТО ВЫ ПОЛУЧАЕТЕ ──────────────────────────────────────────────────────
heading(doc, "Что вы получаете")

callout_box(doc,
    "Это не утилита для сокращения рутины. Это инструмент снижения операционного риска, "
    "ускорения оборота и масштабирования без роста штата."
)

add_table(doc,
    headers=["", "Сейчас", "После внедрения"],
    rows=[
        ["Время на спецификацию",                  "2–15 часов",    "30 секунд"],
        ["Риск ошибки при переносе данных",         "Есть",          "Исключён"],
        ["Зависимость от конкретного сотрудника",   "Высокая",       "Любой сотрудник"],
        ["Единый формат по всем объектам",          "Нет",           "Да"],
        ["Параллельная обработка объектов",         "Невозможна",    "Без ограничений"],
        ["Скорость реакции на запрос клиента",      "Часы / дни",    "Минуты"],
    ],
    col_widths=[8, 3.5, 4.5],
)

body(doc,
    "При потоке 20–30 объектов в месяц система экономит от 40 до 200 часов "
    "инженерного времени ежемесячно — без найма дополнительных людей."
)

# ── 4. СТОИМОСТЬ И СРОКИ ──────────────────────────────────────────────────────
heading(doc, "Стоимость и сроки")

stage_block(doc,
    title    = "Этап 1 — Анализ, разработка и настройка",
    duration = "3 недели",
    price    = "250 000 ₽",
    items    = [
        "Глубокий анализ ваших DWG-файлов: слои, маркировки, типы панелей, крайние случаи",
        "Разработка и настройка правил распознавания под номенклатуру TechnoPan",
        "Адаптация шаблона Excel под корпоративный стиль",
        "Покрытие нестандартных ситуаций: смешанные форматы, нетиповые объекты",
        "Тестирование на реальных объектах из вашего архива с верификацией результатов",
    ]
)

stage_block(doc,
    title    = "Этап 2 — Внедрение, интеграция и обучение",
    duration = "2 недели",
    price    = "150 000 ₽",
    items    = [
        "Установка и настройка на рабочих местах проектировщиков",
        "Интеграция в текущий рабочий процесс",
        "Обучение сотрудников: инструкция + живые сессии",
        "Сдача с проверкой на реальных текущих объектах",
        "Передача документации и исходных конфигураций",
    ]
)

# итого
total_table = doc.add_table(rows=1, cols=3)
total_table.style = "Table Grid"
data = [("Итого за проект", "Срок: 5 недель", "400 000 ₽")]
colors = [BLUE, BLUE, BLUE]
for i, (val, clr) in enumerate(zip(data[0], colors)):
    cell = total_table.cell(0, i)
    set_cell_bg(cell, clr)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(7)
    add_run(p, val, bold=True, size=13, color=WHITE)
total_table.cell(0, 0).width = Cm(7)
total_table.cell(0, 1).width = Cm(5)
total_table.cell(0, 2).width = Cm(4)

p_pay = doc.add_paragraph()
p_pay.paragraph_format.space_before = Pt(6)
p_pay.paragraph_format.space_after  = Pt(8)
add_run(p_pay, "Оплата: 50% при старте, 50% при сдаче.", size=10, color=GRAY)

# поддержка
heading(doc, "Поддержка после сдачи", level=2)

add_table(doc,
    headers=["Период", "Стоимость"],
    rows=[
        ["Первые 3 месяца", "Бесплатно"],
        ["С 4-го месяца",   "25 000 ₽/мес"],
    ],
    col_widths=[10, 6],
)

body(doc,
    "Поддержка включает доработки под изменение номенклатуры, обновления шаблонов, "
    "консультации сотрудников и адаптацию под новые форматы файлов."
)

# окупаемость
heading(doc, "Окупаемость", level=2)
body(doc, "Если час инженера стоит 800–1 200 ₽, а система экономит 80–200 ч/мес:")
bullet(doc, "Экономия в месяц: 64 000 – 240 000 ₽")
bullet(doc, "Окупаемость проекта: 2–6 месяцев")
body(doc,
    "И это только прямая экономия времени — без учёта ускорения сделок и снижения потерь от ошибок.",
    color=GRAY, size=10
)

# ── 5. ЧТО БУДЕТ ДАЛЬШЕ ──────────────────────────────────────────────────────
heading(doc, "Что будет дальше")

callout_box(doc,
    "Вопрос не в том, автоматизировать или нет. Вопрос — сколько объектов TechnoPan планирует "
    "обрабатывать через 2–3 года, и выдержит ли текущий ручной процесс этот рост "
    "без пропорционального увеличения штата."
)

body(doc, "Автоматизация спецификаций — первый шаг. Дорожная карта расширения:")

roadmap = [
    ("Коммерческий отдел (3–6 мес.)", [
        "Автогенерация КП в Word/PDF из данных спецификации",
        "Расчёт стоимости с учётом прайс-листа, скидок и условий поставки",
        "Готовый документ к отправке клиенту — без участия проектировщика",
    ]),
    ("Производство и склад (6–9 мес.)", [
        "Передача спецификации напрямую в производственный план",
        "Контроль остатков металла, плёнки, комплектующих",
        "Автоматические заявки на закупку при дефиците",
    ]),
    ("Логистика (9–12 мес.)", [
        "Расчёт объёма и веса заказа для транспортировки",
        "Формирование упаковочных и маршрутных листов",
        "Интеграция с транспортными компаниями",
    ]),
    ("Бухгалтерия и документооборот", [
        "Счета, накладные, акты — автоматически из спецификации",
        "Передача данных в 1С",
        "Единый архив проектов и документов",
    ]),
    ("Аналитика для руководителя", [
        "Объём и площадь по заказам в разрезе периодов",
        "Статистика по типам панелей, клиентам, объектам",
        "Воронка: от входящего DWG до выставленного счёта",
    ]),
]

for section_title, items in roadmap:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    add_run(p, section_title, bold=True, size=11, color=ACCENT)
    for item in items:
        bullet(doc, item)

# ── 6. СЛЕДУЮЩИЙ ШАГ ─────────────────────────────────────────────────────────
heading(doc, "Следующий шаг")

callout_box(doc,
    "Мы уже разобрали ваши файлы. Готовы провести демо прямо сейчас — показать, "
    "как система обрабатывает «Склад готовой продукции» и выдаёт готовую спецификацию за 30 секунд."
)

body(doc, "После демо — договор, предоплата 50%, старт через 2 рабочих дня.")

p_contact = doc.add_paragraph()
p_contact.paragraph_format.space_before = Pt(8)
add_run(p_contact, "xteam.pro", bold=True, size=13, color=BLUE)

# ── ПОДВАЛ ────────────────────────────────────────────────────────────────────
doc.add_paragraph().paragraph_format.space_after = Pt(12)
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(footer_p, "XTeam.Pro — автоматизация, которая позволяет расти без роста затрат.",
        italic=True, size=10, color=GRAY)

doc.save(OUT)
print(f"Saved: {OUT}")
