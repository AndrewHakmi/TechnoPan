
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT_FILE = r"c:\Users\freak\TRAE\TechnoPan\Договор_разработки_TechnoPan.md"
OUTPUT_FILE = r"c:\Users\freak\TRAE\TechnoPan\Договор_разработки_TechnoPan.docx"

def convert():
    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        # Headers
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(14)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.bold = True
            run.font.size = Pt(13)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(12)
        
        # Lists
        elif line.startswith("*   ") or line.startswith("-   ") or line.startswith("* ") or line.startswith("- "):
            clean_line = re.sub(r"^[\*\-]\s+", "", line)
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold inside list items
            parts = re.split(r"(\*\*.*?\*\*)", clean_line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
        
        # Horizontal rule
        elif line.startswith("---"):
            doc.add_page_break()
            
        # Normal text with formatting
        else:
            if "ПОДПИСИ СТОРОН" in line:
                 p = doc.add_paragraph()
                 run = p.add_run(line)
                 run.bold = True
                 p.paragraph_format.space_before = Pt(20)
                 continue

            # Check alignment for city/date
            if "г. Новосибирск" in line:
                p = doc.add_paragraph(line)
                # Simple heuristic for this specific line
                continue
                
            p = doc.add_paragraph()
            # Handle bold text **...**
            parts = re.split(r"(\*\*.*?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)

    doc.save(OUTPUT_FILE)
    print(f"Saved to {OUTPUT_FILE}")

if __name__ == "__main__":
    convert()
